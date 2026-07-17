import numpy as np
import docx
import re
import openpyxl
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from doctr.io import DocumentFile
from doctr.models import ocr_predictor
from openpyxl import load_workbook
from openpyxl.styles import NamedStyle, Font, Alignment
import fitz #PyMuPDF
import cv2
from datetime import datetime
from openpyxl import Workbook

#--- OBTENER EL NUMERO DE LA GUIA ---
guia=[]
nombreGuia="guia"
mes="junio"
coordenada={'guia':[925,2810,1512,2900]}

#nombrePDF="Abril2026-0" #Nombre del PDF
# Abrir PDF Guia
docGuia = fitz.open("docs\\Archivos\\guia.pdf") # Lectura del PDF
page = docGuia[0] # Seleccion de la hoja del archivo PDF
zoom = 2  # Renderizar como imagen (controlas resolución aquí)
mat = fitz.Matrix(zoom, zoom)
pix = page.get_pixmap(matrix=mat,dpi=600)
# Convertir a una imagen modelo RGB (3 Colores)
img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
#PREPROCESADO
img_np=np.asarray(img) #Conversion de la imagen a un arreglo bidimencional 3D
# Coordenadas en imagen (ejemplo)

#PREPROCESADO
#img_npGuia=np.asarray(imgGuia) #Conversion de la imagen a un arreglo bidimencional 3D
# Coordenadas en imagen (ejemplo)

img_delimitada=img.crop(coordenada['guia']) #Recorte de la imagen a la seccion Nombre
guia=np.asarray(img_delimitada) #Conversion a matriz para el reconocimiento
#display(img_delimitada) #Desplegar imagen

num_guia=[]
modelo=ocr_predictor(det_arch='db_resnet50',reco_arch='crnn_vgg16_bn',pretrained=True,assume_straight_pages=True)#Modelo para la deteccion y reconocimiento de texto

result=modelo([guia]) #Modelo para la deteccio del numero de guia
json=result.export() #Exportacion del resultado
for pagina in json['pages']:
  for bloque in pagina['blocks']:
    for linea in bloque['lines']:
      linea=([w['value'] for w in linea['words']])
      for x in linea:
        num_guia.append(x)
num_guia="".join(num_guia)
num_guia=re.sub("-","",num_guia)

def date():
  date=datetime.now()
  dia=date.day
  mes=date.month
  anio=date.year
  if(mes==1):
    mes="enero"
  if(mes==2):
    mes="febrero"
  if(mes==3):
    mes="marzo"
  if(mes==4):
    mes="abril"
  if(mes==5):
    mes="mayo"
  if(mes==6):
    mes="junio"
  if(mes==7):
    mes="julio"
  if(mes==8):
    mes="agosto"
  if(mes==9):
    mes="septiembre"
  if(mes==10):
    mes="octubre"
  if(mes==11):
    mes="noviembre"
  if(mes==12):
    mes="diciembre"
  return f"{dia} de {mes} del {anio}"

#FORMATO PARA LA GENERACION DEL ARCHIVO WORD
documento=Document("docs\\Formato\\Acuse.docx") #Lectura del formato word del Acuse
#FECHA
fecha=documento.add_paragraph()
style_fecha=fecha.add_run(f"Tijuana, Baja California a {date()}")
style_fecha.font.size=Pt(11)
style_fecha.font.name="Calibri"
fecha.alignment=WD_ALIGN_PARAGRAPH.CENTER
#PARRAFO
descripcion=documento.add_paragraph()
style_desc=fecha.add_run("Se hace entrega los siguientes Formatos únicos de recepción y resguardos de equipos y componentes, así como de listas de asistencia del personal que se encuentra en sede, para firma del Coordinador Administrativo y se entrega un juego en original que corresponden a los siguientes datos.")
style_desc.font.size=Pt(11)
style_desc.font.name="Calibri"
descripcion.alignment=WD_ALIGN_PARAGRAPH.LEFT
#GUIA
parrafo=documento.add_paragraph()
estilo=parrafo.add_run(f"Numero de guía: {num_guia}")
estilo.font.size=Pt(18)
estilo.font.name="Calibri"
estilo.font.bold=True
#Creer tabla 3x4
tabla=documento.add_table(rows=3,cols=4,style="Table Grid")
#Encabezado de la tabla
hdr=tabla.rows[0].cells
hdr[0].text="Folio"
hdr[1].text="Usuario"
hdr[2].text="Ubicación"
hdr[3].text="Serie Principal del Equipo"
#Estilo de encabezado
for cell in hdr:
  parrafo = cell.paragraphs[0]
  for run in parrafo.runs:
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
#Creacion de renglones 2 y 3
lst1=tabla.rows[1].cells
lst1[0].text="LISTA DE ASISTENCIA"
lst1[1].text="ALEX RIVERA PEREZ"
lst1[2].text="SEDE TIJUANA"
lst2=tabla.rows[2].cells
lst2[0].text="LISTA DE ASISTENCIA"
lst2[1].text="ERIC ARMANDO HERRERA PARRA"
lst2[2].text="SEDE TIJUANA"

#Estilo de renglones 2 y 3
for cell in lst1:
  parrafo=cell.paragraphs[0]
  for run in parrafo.runs:
    run.font.size=Pt(11)
    run.font.name="Calibri"
for cell in lst2:
  parrafo=cell.paragraphs[0]
  for run in parrafo.runs:
    run.font.size=Pt(11)
    run.font.name="Calibri"

#Alineacion de texto centrado a toda la tabla
for renglon in tabla.rows:
  for celda in renglon.cells:
    for parraf in celda.paragraphs:
      parraf.alignment=WD_ALIGN_PARAGRAPH.CENTER

datos_acta={}
datos_resguardo={}
cantidad=0
can_res=0
coordenadas={'ticket':[300,500,3000,750],'datos':[300,750,4700,2250]}  #ACTA: [500,500,3000,750]
resultado=[]

#DEFINICION DEL METODO PARA LA CREACION DE LA ACTA
def acta(mes,acta):     #acta(mes,"acta-1")
  global cantidad
  nuevo=[]
  ticket=reconocimiento_acta(mes,acta,"ticket")
  datos=reconocimiento_acta(mes,acta,"datos")
  if(len(ticket)==1):
    if(re.search("-",ticket[0])):
      tickets=re.split("-",ticket[0])
      nuevo.append(tickets[0])
      nuevo.append(tickets[1])
  if(len(datos)==3):
    for x in range(len(datos)):
      nuevo.append(datos[x][0])
  if(len(nuevo)==5):
    #def nuevo_renglon(folio,usuario,ubicacion,serie):
    nuevo_renglon((nuevo[0]+"-"+nuevo[1]),nuevo[2],nuevo[3],nuevo[4])
    #Llenar arreglo datos
    datos_acta[cantidad]={
        'FGR':nuevo[0],
        'THEOS':nuevo[1],
        'USUARIO':nuevo[2],
        'CEDE':nuevo[3],
        'SERIE':nuevo[4],
        'GUIA':num_guia
    }
  cantidad+=1
  #print(cantidad)
  #Agregar excel Acta
  Excel_Actas(datos_acta)

  #guardar cambios al docx
  documento.save(f"docs/Nuevo/ACUSE BC,{mes}2026.docx")

  return datos_acta

def reconocimiento_acta(mes,acta,tipo): # ("abril","acta-1","ticket")
  arreglo=[]
  v_nombre=[]
  v_cede=[]
  v_serie=[]
  v_theos=""
  v_fgr=""
  #Lee imagen del acta
  doc = fitz.open(f"docs/Archivos/{acta}.pdf") # Lectura del PDF
  modelo=ocr_predictor(det_arch='db_resnet50',reco_arch='crnn_vgg16_bn',pretrained=True,assume_straight_pages=True)
  page = doc[0] # Seleccion de la hoja del archivo PDF
  zoom = 2  # Renderizar como imagen (controlas resolución aquí)
  mat = fitz.Matrix(zoom, zoom)
  pix = page.get_pixmap(matrix=mat,dpi=600)
  # Convertir a una imagen modelo RGB (3 Colores)
  img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
  # utiliza el metodo crop para delimitar la imagen (izquierda, arriba, derecha, abajo).
  img_delimitada = img.crop(coordenadas[tipo])
  # convierte la imagen PIL a una arreglo numpy para poder usarlo con el modelo ocr
  delimitado = np.asarray(img_delimitada)
  result = modelo([delimitado])
  json = result.export()
  #display(img_delimitada)
  for pagina in json['pages']:
    for bloque in pagina['blocks']:
      for linea in bloque['lines']:
        linea=" ".join([w['value'] for w in linea['words']])
        arreglo.append(linea)
  if(tipo=="ticket"):
    for x in range(len(arreglo)): #['NO. DE TICKET FGR:', 'INC25368', 'NO. DE TICKET THEOS:', 'INC8942']
      if(re.search("FGR",arreglo[x])):
        v_fgr=arreglo[x+1]
      if(re.search("THEOS",arreglo[x])):
        v_theos=arreglo[x+1]
    if(v_fgr!="" and v_theos!=""):
      resultado=v_fgr+"-"+v_theos
      resultado=[resultado]
  if(tipo=="datos"):
    for x in range(len(arreglo)):
      if(re.search("USUARIO",arreglo[x])):
        if(re.search(":",arreglo[x])):
          nuevo=re.split(":",arreglo[x])
          v_nombre.append(nuevo[1])
      if(re.search("SEDE",arreglo[x])):
        v_cede.append(arreglo[x])
      if(re.search("SERIE",arreglo[x])):
        v_serie.append(arreglo[x+1])
      resultado=[v_nombre,v_cede,v_serie]
  return resultado

#Metodo para agregar nuevo renglon a la tabla
def nuevo_renglon(folio,usuario,ubicacion,serie): #def nuevo_renglon(folio,usuario,ubicacion,serie):
  renglon=tabla.add_row().cells
  renglon[0].text=folio
  renglon[1].text=usuario
  renglon[2].text=ubicacion
  renglon[3].text=serie
  #Formato del renglon
  for cell in renglon:
    parrafo=cell.paragraphs[0]
    for run in parrafo.runs:
      run.font.size=Pt(11)
      run.font.name="Calibri"
    parrafo.alignment=WD_ALIGN_PARAGRAPH.CENTER

def reconocimiento_resguardo(mes,resguardo):
  arreglo=[]
  #Lee imagen del acta

  coordenadas={'datos':[0,0,5100,3000],'ticket':[0,5800,5080,6250]}

  def reconocimiento(coordenada):
    doc = fitz.open(f"docs/Archivos/{resguardo}.pdf") # Lectura del PDF
    modelo=ocr_predictor(det_arch='db_resnet50',reco_arch='crnn_vgg16_bn',pretrained=True,assume_straight_pages=True)
    page = doc[0] # Seleccion de la hoja del archivo PDF
    zoom = 2  # Renderizar como imagen (controlas resolución aquí)
    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat,dpi=600)
    # Convertir a una imagen modelo RGB (3 Colores)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    # utiliza el metodo crop para delimitar la imagen (izquierda, arriba, derecha, abajo).

    img_delimitada = img.crop(coordenada)
    # convierte la imagen PIL a una arreglo numpy para poder usarlo con el modelo ocr
    delimitado = np.asarray(img_delimitada)
    result = modelo([delimitado])
    json = result.export()
    #display(img_delimitada)

    for pagina in json['pages']:
      for bloque in pagina['blocks']:
        for linea in bloque['lines']:
          linea=" ".join([w['value'] for w in linea['words']])
          arreglo.append(linea)
    return arreglo

  reconocimiento(coordenadas['datos'])
  datos=reconocimiento(coordenadas['ticket'])
  return datos

#Genera un nuevo resguardo -----------------------------------------------------
def resguardo(mes,archivo): #(mes,"resguardo-1") mes="junio"
  folio=[]
  id=[]
  serie=[]
  estado=[]
  usuario=[]
  #arr_datos=[]
  datos=reconocimiento_resguardo(mes,archivo) #(mes,"resguardo-1") mes="junio"
  global can_res
  for x in range(len(datos)):
    if(re.search("FOLIO",datos[x])): # ID, USUARIO, SERIE, ESTADO, TICKET, GUIA
      if(re.search(':',datos[x])):
        folio=re.split(':',datos[x])
        folio=folio[1]
        folio=re.sub(" ","",folio)
    if(re.search("ID:",datos[x])):
      id=re.split(":",datos[x])
      id=id[1]
      id=re.sub(" ","",id)
    if(re.search("SERIE",datos[x])):
      serie=datos[x+6]
    if(re.search("ESTADO:",datos[x])):
      estado=re.split(":",datos[x])
      estado=estado[1]
    if(re.search("INMUEBLE",datos[x])): #'INMUEBLE: SEDE TIJUANA'
      if(re.search(":",datos[x])):
        inmueble=re.split(":",datos[x])
        inmueble=inmueble[1]
        txt=""
        n=""
        if(inmueble[0]==" "):
          for x in range(1,len(inmueble)):
            txt=txt+inmueble[x]
        if(txt[len(txt)-1]==" "):
          for x in range(0,len(txt)-1):
            n=n+txt[x]
            inmueble=n
        if(inmueble!="SEDE TIJUANA" or inmueble!="SUBSEDE TIJUANA" or inmueble!="SUBSEDE ENSENADA" or inmueble!="SUBSEDE MEXICALI"):
          inmueble="INVALIDO"
    #'NOMBRE(S): DANIELA','APELLIDO PATERNO: ENRIQUEZ', 'APELLIDO MATERNO: CAMINO'
    if(re.search("NOMBRE",datos[x])):
      if(re.search(":",datos[x])):
        nombre=re.split(":",datos[x])
        usuario.append(nombre[1])
    if(re.search("PATERNO",datos[x])):
      if(re.search(":",datos[x])):
        apellido=re.split(":",datos[x])
        usuario.append(apellido[1])
    if(re.search("MATERNO",datos[x])):
      if(re.search(":",datos[x])):
        apellido=re.split(":",datos[x])
        usuario.append(apellido[1])
    if(re.search("TICKET",datos[x])):
      if(re.search(":",datos[x])):
        ticket=re.split(":",datos[x])
        ticket=ticket[1]
  if(len(usuario)==3):
    usuario="".join(usuario)

  datos_resguardo[can_res]={
      'FOLIO':folio,
      'ID':id,
      'SERIE':serie,
      'ESTADO':estado,
      'INMUEBLE':inmueble,
      'USUARIO':usuario,
      'TICKET': ticket
  }

  #Nuevo renglon docx
  nuevo_renglon(folio,usuario,inmueble,serie) #nuevo_renglon(folio,usuario,ubicacion,serie)

  #Generar excel de resguardo
  Excel_Resguardos(datos_resguardo[can_res])

  #Guardar cambios al docx
  documento.save(f"docs/Nuevo/ACUSE BC,{mes}2026.docx")
  print("ACUSE Creado")
  can_res+=1

  return datos_resguardo

from openpyxl import workbook
#Estilo del titulo
estiloEncabezado=NamedStyle(name="EstiloEncabezado")
estiloEncabezado.font=Font(bold=True,size=12,color="000000")
estiloEncabezado.alignment=Alignment(horizontal="center", vertical="center",wrapText=True)

#Estilo de los renglones
estiloRenglon=NamedStyle(name="EstiloRenglon")
estiloRenglon.font=Font(size=12,color="000000")
estiloRenglon.alignment=Alignment(horizontal="center",vertical="center")

def Excel_Resguardos(datos):
  global estiloEncabezado,estiloRenglon
  ren=2
  cantidad=len(datos_resguardo) #NO.|FOLIO|ID|NOMBRE DEL USUARIO|SERIE|ESTADO|INMUEBLE|TICKET|GUIA
  #Cargar Excel
  #archivo=load_workbook("drive/MyDrive/PyFormatos/Formato/Formato_Excel.xlsx")
  archivo=Workbook()
  hoja=archivo.active
  #Tamaño de celdas
  hoja.column_dimensions['A'].width = 7
  hoja.column_dimensions['B'].width = 15
  hoja.column_dimensions['C'].width = 50
  hoja.column_dimensions['D'].width = 20
  hoja.column_dimensions['E'].width = 20
  hoja.column_dimensions['F'].width = 30
  hoja.column_dimensions['G'].width = 20

  archivo.add_named_style(estiloEncabezado)
  archivo.add_named_style(estiloRenglon)

  celda=hoja.cell(row=1,column=1,value="NO.")
  celda.style="EstiloEncabezado"
  celda=hoja.cell(row=1,column=2,value="ID")
  celda.style="EstiloEncabezado"
  celda=hoja.cell(row=1,column=3,value="NOMBRE DEL USUARIO")
  celda.style="EstiloEncabezado"
  celda=hoja.cell(row=1,column=4,value="SERIE")
  celda.style="EstiloEncabezado"
  celda=hoja.cell(row=1,column=5,value="ESTADO")
  celda.style="EstiloEncabezado"
  celda=hoja.cell(row=1,column=6,value="TICKET")
  celda.style="EstiloEncabezado"
  celda=hoja.cell(row=1,column=7,value="GUIA")
  celda.style="EstiloEncabezado"

  for x in range(cantidad):
    celda=hoja.cell(row=ren,column=1,value=x)
    celda.style="EstiloRenglon"
    celda=hoja.cell(row=ren,column=2,value=datos_resguardo[x]['ID'])
    celda.style="EstiloRenglon"
    celda=hoja.cell(row=ren,column=3,value=datos_resguardo[x]['USUARIO'])
    celda.style="EstiloRenglon"
    celda=hoja.cell(row=ren,column=4,value=datos_resguardo[x]['SERIE'])
    celda.style="EstiloRenglon"
    celda=hoja.cell(row=ren,column=5,value=datos_resguardo[x]['ESTADO'])
    celda.style="EstiloRenglon"
    celda=hoja.cell(row=ren,column=6,value=datos_resguardo[x]['TICKET'])
    celda.style="EstiloRenglon"
    celda=hoja.cell(row=ren,column=7,value=num_guia)
    celda.style="EstiloRenglon"
    ren+=1
  archivo.save(f"docs/Nuevo/Resguardos{mes}2026.xlsx")
  print("Resguardo Excel Creado")

def Excel_Actas(datos):
  global estiloRenglon,estiloEncabezado
  cantidad=len(datos_acta)
  ren=2
  #LLENAR EXCEL
  archivo=Workbook()
  hoja=archivo.active
  hoja.column_dimensions['A'].width = 10
  hoja.column_dimensions['B'].width = 20
  hoja.column_dimensions['C'].width = 20
  hoja.column_dimensions['D'].width = 20
  
  archivo.add_named_style(estiloEncabezado)
  archivo.add_named_style(estiloRenglon)

  celda=hoja.cell(row=1,column=1,value="NO.")
  celda.style="EstiloEncabezado"
  celda=hoja.cell(row=1,column=2,value="NO. DE TICKET FGR")
  celda.style="EstiloEncabezado"
  celda=hoja.cell(row=1,column=3,value="NO. DE TICKET THEOS")
  celda.style="EstiloEncabezado"
  celda=hoja.cell(row=1,column=4,value="GUIA")
  celda.style="EstiloEncabezado"

  for x in range(cantidad):
    celda=hoja.cell(row=ren,column=1,value=x+1)
    celda.style="EstiloRenglon"
    celda=hoja.cell(row=ren,column=2,value=datos_acta[x]['FGR'])
    celda.style="EstiloRenglon"
    celda=hoja.cell(row=ren,column=3,value=datos_acta[x]['THEOS'])
    celda.style="EstiloRenglon"
    celda=hoja.cell(row=ren,column=4,value=num_guia)
    celda.style="EstiloRenglon"
    ren+=1

  archivo.save(f"docs/Nuevo/Actas{mes}2026.xlsx")
  print("Actas Excel Creado")

#Nombre de la imagen del acta (5 min)

#print("--Acta--")
acta(mes,"acta-1")
print(acta(mes,"acta-2"))

#print("--Resguardo--")
resguardo(mes,"resguardo-1")
resguardo(mes,"resguardo-2")
print(resguardo(mes,"resguardo-3"))